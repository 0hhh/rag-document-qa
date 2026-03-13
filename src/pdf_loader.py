# from pypdf import PdfReader
# from langchain_core.documents import Document
# def load_pdf(file):
#     reader = PdfReader(file)
#     documents = []
#     for page_number, page in enumerate(reader.pages):
#         text = page.extract_text()
#         if text:
#             documents.append(
#                 Document(
#                     page_content=text,
#                     metadata={"page": page_number + 1,
#                               "source": file.name}
#                 )
#             )
#     return documents
from pypdf import PdfReader
from langchain_core.documents import Document
import docx
import pandas as pd


def load_pdf(file):

    name = file.name.lower()
    documents = []

    # -------------------
    # PDF
    # -------------------
    if name.endswith(".pdf"):

        reader = PdfReader(file)

        for page_number, page in enumerate(reader.pages):

            text = page.extract_text()

            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "page": page_number + 1,
                            "source": file.name
                        }
                    )
                )

    # -------------------
    # DOCX
    # -------------------
    elif name.endswith(".docx"):

        doc = docx.Document(file)

        text = "\n".join([para.text for para in doc.paragraphs])

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "page": 1,
                    "source": file.name
                }
            )
        )

    # -------------------
    # TXT
    # -------------------
    elif name.endswith(".txt"):

        text = file.read().decode("utf-8")

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "page": 1,
                    "source": file.name
                }
            )
        )

    # -------------------
    # CSV
    # -------------------
    elif name.endswith(".csv"):

        df = pd.read_csv(file)

        # text = df.to_string()
        text = df.to_markdown()
        documents.append(
            Document(
                page_content=text,
                metadata={
                    "page": 1,
                    "source": file.name
                }
            )
        )

    else:
        raise ValueError("Unsupported file format")

    return documents